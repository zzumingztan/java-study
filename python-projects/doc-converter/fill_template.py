import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT

def set_font_style(run, font_name='SimSun', font_size=Pt(10.5)):
    """设置字体样式（五号字体=10.5pt）"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size

def fill_template_with_content(template_path, md_content, output_path):
    # 读取模板文档
    doc = Document(template_path)
    
    # 提取Markdown内容中的各个部分
    lines = md_content.split('\n')
    
    # 提取各部分内容
    sections = {}
    current_section = None
    
    for line in lines:
        line = line.strip()
        if line.startswith('## 一、实验目的'):
            current_section = 'purpose'
            sections[current_section] = []
        elif line.startswith('## 二、实验内容'):
            current_section = 'content'
            sections[current_section] = []
        elif line.startswith('## 三、实验环境'):
            current_section = 'environment'
            sections[current_section] = []
        elif line.startswith('## 四、实验步骤'):
            current_section = 'steps'
            sections[current_section] = []
        elif line.startswith('## 五、实验结果与分析'):
            current_section = 'results'
            sections[current_section] = []
        elif line.startswith('## 六、实验总结'):
            current_section = 'summary'
            sections[current_section] = []
        elif current_section and line and not line.startswith('## ') and not line.startswith('---'):
            # 过滤掉代码块、图片等
            if not line.startswith('```') and not line.startswith('![') and not line.startswith('['):
                sections[current_section].append(line)
    
    # 查找并替换占位符
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == '实验目的' or text == '实验内容' or text == '实验环境' or \
           text == '实验步骤' or text == '实验结果与分析' or text == '实验总结':
            # 找到对应的占位符段落
            next_para = None
            for i, p in enumerate(doc.paragraphs):
                if p == paragraph:
                    if i + 1 < len(doc.paragraphs):
                        next_para = doc.paragraphs[i + 1]
                    break
            
            if next_para and '（具体内容' in next_para.text:
                # 清空占位符段落
                next_para.clear()
                
                # 添加对应的内容
                section_key = None
                if text == '实验目的':
                    section_key = 'purpose'
                elif text == '实验内容':
                    section_key = 'content'
                elif text == '实验环境':
                    section_key = 'environment'
                elif text == '实验步骤':
                    section_key = 'steps'
                elif text == '实验结果与分析':
                    section_key = 'results'
                elif text == '实验总结':
                    section_key = 'summary'
                
                if section_key in sections and sections[section_key]:
                    # 添加内容
                    for line in sections[section_key]:
                        if line.strip() and not line.strip().startswith('#') and not line.strip().startswith('*'):
                            p = doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Inches(0.28)
                            p.paragraph_format.line_spacing = 1.0
                            for run in p.runs:
                                set_font_style(run)
    
    # 保存填充后的文档
    doc.save(output_path)

# 读取Markdown文件
try:
    with open('房价预测实验报告.md', 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 填充模板（使用你提供的模板）
    fill_template_with_content('c:/Users/21268/Desktop/人工智能导论/人工智能导论实验报告模板.docx', 
                              md_content, 
                              '房价预测实验报告_填充版.docx')
    print("模板填充完成！DOCX文件已生成：房价预测实验报告_填充版.docx")
except Exception as e:
    print(f"模板填充失败：{e}")