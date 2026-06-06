import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font_style(run, font_name='SimSun', size=12, bold=False, italic=False):
    """辅助函数：设置运行块的字体（宋体）和西文字体（Times New Roman）"""
    run.font.name = 'Times New Roman'  # 西文字体
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)  # 中文字体
    run.font.size = Pt(size)
    run.bold = bold
    run.font.italic = italic

def create_report_docx(output_path):
    doc = Document()

    # --- 1. 设置全局页边距 (A4 标准) ---
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # --- 2. 封面：学院名称 ---
    p1 = doc.add_paragraph()
    p1.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run1 = p1.add_run("郑州大学计算机与人工智能学院")
    set_font_style(run1, font_name='Microsoft YaHei', size=22, bold=True) # 使用雅黑更接近标题效果

    # --- 3. 封面：大标题 ---
    p2 = doc.add_paragraph()
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run2 = p2.add_run("人工智能实验报告")
    set_font_style(run2, font_name='SimSun', size=20, bold=True)
    # 给标题加下划线（模拟模板效果）
    run2.underline = True

    doc.add_paragraph() # 空行

    # --- 4. 封面：实验题目 ---
    p3 = doc.add_paragraph()
    p3.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run3 = p3.add_run("实验题目：基于全连接网络实现房价预测")
    set_font_style(run3, font_name='SimSun', size=14, bold=True)

    doc.add_paragraph() # 空行

    # --- 5. 封面：信息表格 (核心修改) ---
    # 创建 5 行 2 列的表格
    table = doc.add_table(rows=5, cols=2)
    table.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # 定义要填入的信息
    info_data = [
        ("专业班级", "人工智能六班"),
        ("学号", "202415060613"),
        ("姓名", "李润豪"),
        ("指导老师", "王可"),
        ("报告日期", "2026 年 05 月")
    ]

    for i, (label, value) in enumerate(info_data):
        # 左侧单元格：标签（右对齐）
        cell_l = table.cell(i, 0)
        cell_l.text = f"**{label}**" if label else "" 
        p_l = cell_l.paragraphs[0]
        p_l.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        set_font_style(p_l.runs[0], font_name='SimSun', size=12)

        # 右侧单元格：内容（下划线效果）
        cell_r = table.cell(i, 1)
        p_r = cell_r.paragraphs[0]
        p_r.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run_r = p_r.add_run(value)
        set_font_style(run_r, font_name='SimSun', size=12)
        # 给值加边框/下划线（通过直接修改 XML 实现或简单用文本表达）

    doc.add_page_break() # 换页

    # --- 6. 正文部分 (按照截图的一、二、三... 结构) ---
    sections_content = [
        ("一、实验目的", "加深对神经网络、回归等概念的理解，掌握模型的构建、损失函数的定义、优化器的选择等关键操作。"),
        ("二、实验内容", "使用飞桨搭建单层的全连接网络，利用 uci-housing 数据集训练模型并保存参数。"),
        ("三、实验环境", "Windows 10/11, VS Code, PaddlePaddle 2.0.2。"),
        ("四、实验步骤", "1. 数据准备与归一化；2. 定义 Regressor 网络结构；3. 训练模型并绘制 Loss 曲线。"),
        ("五、实验结果与分析", "Loss 曲线从约 800 快速下降并收敛至 100 以下，模型在高价位预测略有偏差。"),
        ("六、实验总结", "掌握了飞桨框架的基本用法，理解了反向传播在参数优化中的作用。")
    ]

    for title, body in sections_content:
        # 添加标题
        p_t = doc.add_paragraph()
        run_t = p_t.add_run(title)
        set_font_style(run_t, font_name='SimSun', size=14, bold=True)
        
        # 添加正文（首行缩进两个字符）
        p_b = doc.add_paragraph()
        p_b.paragraph_format.first_line_indent = Pt(24) # 约等于两个 12pt 字符
        run_b = p_b.add_run(body)
        set_font_style(run_b, font_name='SimSun', size=12)
        doc.add_paragraph() # 段后空行

    # 保存文件
    doc.save(output_path)
    print(f"转换成功！文件已生成：{output_path}")

# 执行生成
create_report_docx('房价预测实验报告_符合格式版.docx')