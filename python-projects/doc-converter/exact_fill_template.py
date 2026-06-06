import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT

def set_font_style(run, font_name='SimSun', font_size=Pt(10.5)):
    """设置字体样式（五号字体=10.5pt）"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size

def extract_sections_from_content(content_doc_path):
    """从内容文档中提取各部分内容"""
    doc = Document(content_doc_path)
    
    # 提取各部分内容
    sections = {
        'purpose': [],
        'content': [],
        'environment': [],
        'steps': [],
        'results': [],
        'summary': []
    }
    
    current_section = None
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text == '、实验目的':
            current_section = 'purpose'
        elif text == '、实验内容':
            current_section = 'content'
        elif text == '、实验环境':
            current_section = 'environment'
        elif text == '、实验步骤':
            current_section = 'steps'
        elif text == '、实验结果与分析':
            current_section = 'results'
        elif text == '、实验总结':
            current_section = 'summary'
        elif current_section and text and not text.startswith('---') and not text.startswith('## ') and not text.startswith('### '):
            # 过滤掉空行和分隔线
            if text != '---' and text != '':
                sections[current_section].append(text)
    
    return sections

def fill_template_exact(template_path, content_sections, output_path):
    """精确填充模板"""
    doc = Document(template_path)
    
    # 查找并替换各部分占位符
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        
        if text == '实验目的':
            # 找到下一个段落（占位符）
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                if '（具体内容' in next_para.text:
                    # 清空占位符段落
                    next_para.clear()
                    # 添加内容
                    for line in content_sections['purpose']:
                        if line.strip():
                            p = doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Inches(0.28)
                            p.paragraph_format.line_spacing = 1.0
                            for run in p.runs:
                                set_font_style(run)
        
        elif text == '实验内容':
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                if '（具体内容' in next_para.text:
                    next_para.clear()
                    for line in content_sections['content']:
                        if line.strip():
                            p = doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Inches(0.28)
                            p.paragraph_format.line_spacing = 1.0
                            for run in p.runs:
                                set_font_style(run)
        
        elif text == '实验环境':
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                if '（具体内容' in next_para.text:
                    next_para.clear()
                    for line in content_sections['environment']:
                        if line.strip():
                            p = doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Inches(0.28)
                            p.paragraph_format.line_spacing = 1.0
                            for run in p.runs:
                                set_font_style(run)
        
        elif text == '实验步骤':
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                if '（具体内容' in next_para.text:
                    next_para.clear()
                    for line in content_sections['steps']:
                        if line.strip():
                            p = doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Inches(0.28)
                            p.paragraph_format.line_spacing = 1.0
                            for run in p.runs:
                                set_font_style(run)
        
        elif text == '实验结果与分析':
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                if '（具体内容' in next_para.text:
                    next_para.clear()
                    for line in content_sections['results']:
                        if line.strip():
                            p = doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Inches(0.28)
                            p.paragraph_format.line_spacing = 1.0
                            for run in p.runs:
                                set_font_style(run)
        
        elif text == '实验总结':
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                if '（具体内容' in next_para.text:
                    next_para.clear()
                    for line in content_sections['summary']:
                        if line.strip():
                            p = doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Inches(0.28)
                            p.paragraph_format.line_spacing = 1.0
                            for run in p.runs:
                                set_font_style(run)
    
    # 保存填充后的文档
    doc.save(output_path)

# 提取内容
try:
    content_sections = extract_sections_from_content('f:/huibian/房价预测实验报告_模板样式.docx')
    
    # 填充模板
    fill_template_exact('c:/Users/21268/Desktop/人工智能导论/人工智能导论实验报告模板.docx', 
                       content_sections, 
                       '房价预测实验报告_最终完美版.docx')
    print("精确填充完成！DOCX文件已生成：房价预测实验报告_最终完美版.docx")
except Exception as e:
    print(f"精确填充失败：{e}")