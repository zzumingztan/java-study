import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font_style(run, font_name='SimSun', font_size=Pt(10.5)):
    """设置字体样式（五号字体=10.5pt）"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size

def fill_template_with_content(template_path, content, output_path):
    """填充模板"""
    doc = Document(template_path)
    
    # 查找并替换各部分占位符
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        
        if text == '实验目的':
            # 找到下一个段落（应该是占位符）
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                if '（具体内容' in next_para.text:
                    # 清空占位符段落
                    next_para.clear()
                    # 添加内容
                    for line in content['purpose']:
                        if line.strip():
                            p = doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Inches(0.28)
                            p.paragraph_format.line_spacing = 1.0
                            for run in p.runs:
                                set_font_style(run)
        
        elif text == '实验内容':
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                if '（具体内容' in next_para.text:
                    next_para.clear()
                    for line in content['content']:
                        if line.strip():
                            p = doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Inches(0.28)
                            p.paragraph_format.line_spacing = 1.0
                            for run in p.runs:
                                set_font_style(run)
        
        elif text == '实验环境':
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                if '（具体内容' in next_para.text:
                    next_para.clear()
                    for line in content['environment']:
                        if line.strip():
                            p = doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Inches(0.28)
                            p.paragraph_format.line_spacing = 1.0
                            for run in p.runs:
                                set_font_style(run)
        
        elif text == '实验步骤':
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                if '（具体内容' in next_para.text:
                    next_para.clear()
                    for line in content['steps']:
                        if line.strip():
                            p = doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Inches(0.28)
                            p.paragraph_format.line_spacing = 1.0
                            for run in p.runs:
                                set_font_style(run)
        
        elif text == '实验结果与分析':
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                if '（具体内容' in next_para.text:
                    next_para.clear()
                    for line in content['results']:
                        if line.strip():
                            p = doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Inches(0.28)
                            p.paragraph_format.line_spacing = 1.0
                            for run in p.runs:
                                set_font_style(run)
        
        elif text == '实验总结':
            if i + 1 < len(doc.paragraphs):
                next_para = doc.paragraphs[i + 1]
                if '（具体内容' in next_para.text:
                    next_para.clear()
                    for line in content['summary']:
                        if line.strip():
                            p = doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Inches(0.28)
                            p.paragraph_format.line_spacing = 1.0
                            for run in p.runs:
                                set_font_style(run)
    
    # 保存
    doc.save(output_path)

# 提供的内容
content = {
    'purpose': [
        "加深对神经网络、回归等概念的理解，掌握模型的构建、损失函数的定义、优化器的选择、网络的前向计算、网络的反向传播等关键操作；并且熟悉使用飞桨（PaddlePaddle）深度学习框架来进行以上的操作。",
        ""
    ],
    'content': [
        "飞桨内部集成了波士顿房价数据集 uci-housing，该数据集共 506 条数据，每条数据包含 14 个值，其中前 13 个值用来描述房屋的各种信息，作为特征信息，最后一个值为该类房屋价格中位数，即要回归预测的值。",
        "本实验要求使用飞桨搭建单层的全连接网络，并利用上述数据集训练该模型；训练结束后保存模型，存储模型的参数；评估网络时，从参数文件中加载模型参数，再评估模型。",
        ""
    ],
    'environment': [
        "- **硬件资源：** GPU",
        "- **软件环境：** Python 3+, PaddlePaddle 2.0.2, NumPy, Matplotlib",
        "- **开发平台：** Baidu AI Studio",
        ""
    ],
    'steps': [
        "### 步骤一：数据加载及预处理",
        "飞桨中提供了读取 uci-housing 训练集和测试集的接口，分别为 `paddle.text.datasets.UCIHousing(mode='train')` 和 `paddle.text.datasets.UCIHousing(mode='test')`。",
        "数据集加载后，使用 DataLoader 类进行封装，指定 batch_size 定制批量迭代的批大小；shuffle 参数控制样本顺序，使得不同类型样本可以随机分布。本实验训练集的 batch_size 设置为 32，测试集的 batch_size 设置为 8。",
        "核心代码如下：",
        "```python",
        "import paddle",
        "import numpy as np",
        "import matplotlib.pyplot as plt",
        "",
        "# 设置paddle默认的全局数据类型为float64",
        "paddle.set_default_dtype('float64')",
        "",
        "# 加载数据",
        "train_dataset = paddle.text.datasets.UCIHousing(mode='train')",
        "eval_dataset = paddle.text.datasets.UCIHousing(mode='test')",
        "",
        "# 封装训练数据",
        "train_loader = paddle.io.DataLoader(train_dataset, batch_size=32, shuffle=True)",
        "eval_loader = paddle.io.DataLoader(eval_dataset, batch_size=8, shuffle=False)",
        "",
        "print(\"步骤一完成：数据已加载并封装完毕！\")",
        "```",
        "",
        "### 步骤二：模型配置",
        "本实验使用 `paddle.nn.Linear(in_features, out_features)` 实现线性变换。其中，in_features 为输入特征的维度（13 个特征），out_features 为输出特征的维度（1 个预测值）。",
        "在自定义网络模型时，继承 `paddle.nn.Layer` 类，该类基于 OOD 实现的动态图，实现了训练模式和验证模式。在 forward 函数中定义网络从前到后的完整计算过程。",
        "全连接网络的框架如下：",
        "```python",
        "import paddle.nn as nn",
        "",
        "# 定义模型架构",
        "class Regressor(paddle.nn.Layer):",
        "def __init__(self):",
        "super(Regressor, self).__init__()",
        "# 定义一层全连接网络，输入13个特征，输出1个预测值，不需要激活函数",
        "self.fc = nn.Linear(in_features=13, out_features=1)",
        "",
        "def forward(self, inputs):",
        "# 前向计算过程",
        "x = self.fc(inputs)",
        "return x",
        "",
        "# 实例化模型",
        "model = Regressor()",
        "paddle.summary(model, (32, 13), dtypes='float64')",
        "print(\"步骤二完成：全连接网络模型架构已搭建并实例化！\")",
        "```",
        "**运行结果：**",
        "(显示模型各层结构及参数量，共 14 个参数：13 个权重 + 1 个偏置)",
        "",
        "### 步骤三：模型训练",
        "定义好模型框架后：",
        "1. 实例化模型 model，在训练阶段调用 `model.train()` 开启训练模式，该模式具备反向传播功能。",
        "2. 定义损失函数：在回归任务中，使用均方误差损失函数 `paddle.nn.MSELoss(reduction='mean')`，返回 loss 的均值。",
        "3. 定义优化器：使用 `paddle.optimizer.SGD(learning_rate=0.0005)` 进行优化，学习率为 0.0005。",
        "4. 将数据分批次送入模型，根据输出结果和 label 计算 loss，然后执行梯度反向传播更新参数。默认迭代 200 轮（epochs=200）。",
        "在模型训练过程中输出训练损失值，每 10 个批次输出一次，并将批次信息和损失信息存入列表，用于后续绘图。",
        "核心代码如下：",
        "```python",
        "Batch = 0",
        "# 记录批次",
        "Batchs = []",
        "# 记录损失值，用于后续绘图",
        "all_train_loss = []",
        "",
        "# 模型实例化",
        "model = Regressor()",
        "# 训练模型",
        "model.train()",
        "# 均方误差损失函数",
        "mse_loss = paddle.nn.MSELoss()",
        "# 优化函数，需要传入学习率和模型参数",
        "opt = paddle.optimizer.SGD(learning_rate=0.0005, parameters=model.parameters())",
        "",
        "# 迭代次数",
        "epochs_num = 200",
        "for pass_num in range(epochs_num):",
        "for batch_id, data in enumerate(train_loader()):",
        "# data包括input, label",
        "image = data[0]",
        "label = data[1]",
        "",
        "# 前向计算",
        "predict = model(image)",
        "loss = mse_loss(predict, label)",
        "",
        "# 10次迭代记录一次损失值",
        "if batch_id != 0 and batch_id % 10 == 0:",
        "Batch = Batch + 10",
        "Batchs.append(Batch)",
        "# 记录损失值并打印",
        "all_train_loss.append(loss.numpy()[0])",
        "print('epoch:{}, step:{}, train_loss:{}'.format(pass_num, batch_id, loss.numpy()[0]))",
        "",
        "# 反向传播",
        "loss.backward()",
        "# 更新参数",
        "opt.step()