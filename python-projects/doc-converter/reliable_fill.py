import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font_style(run, font_name='SimSun', font_size=Pt(10.5)):
    """设置字体样式（五号字体=10.5pt）"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size

def extract_sections_from_docx(doc_path):
    """从DOCX文件中提取各部分内容"""
    doc = Document(doc_path)
    
    # 查找各部分标题的位置
    sections = {}
    current_section = None
    
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text == '、实验目的':
            current_section = 'purpose'
            sections[current_section] = []
        elif text == '、实验内容':
            current_section = 'content'
            sections[current_section] = []
        elif text == '、实验环境':
            current_section = 'environment'
            sections[current_section] = []
        elif text == '、实验步骤':
            current_section = 'steps'
            sections[current_section] = []
        elif text == '、实验结果与分析':
            current_section = 'results'
            sections[current_section] = []
        elif text == '、实验总结':
            current_section = 'summary'
            sections[current_section] = []
        elif current_section and text and not text.startswith('---') and not text.startswith('## ') and not text.startswith('### '):
            # 添加内容，跳过空行和分隔线
            if text.strip() and text.strip() != '---':
                sections[current_section].append(text)
    
    return sections

def fill_template_reliably(template_path, content_sections, output_path):
    """可靠地填充模板"""
    template_doc = Document(template_path)
    
    # 查找模板中的占位符并替换
    for i, paragraph in enumerate(template_doc.paragraphs):
        text = paragraph.text.strip()
        
        if text == '实验目的':
            # 找到下一个段落（应该是占位符）
            if i + 1 < len(template_doc.paragraphs):
                next_para = template_doc.paragraphs[i + 1]
                if '（具体内容' in next_para.text:
                    # 清空占位符段落
                    next_para.clear()
                    # 添加内容
                    for line in content_sections.get('purpose', []):
                        if line.strip():
                            p = template_doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Inches(0.28)
                            p.paragraph_format.line_spacing = 1.0
                            for run in p.runs:
                                set_font_style(run)
        
        elif text == '实验内容':
            if i + 1 < len(template_doc.paragraphs):
                next_para = template_doc.paragraphs[i + 1]
                if '（具体内容' in next_para.text:
                    next_para.clear()
                    for line in content_sections.get('content', []):
                        if line.strip():
                            p = template_doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Inches(0.28)
                            p.paragraph_format.line_spacing = 1.0
                            for run in p.runs:
                                set_font_style(run)
        
        elif text == '实验环境':
            if i + 1 < len(template_doc.paragraphs):
                next_para = template_doc.paragraphs[i + 1]
                if '（具体内容' in next_para.text:
                    next_para.clear()
                    for line in content_sections.get('environment', []):
                        if line.strip():
                            p = template_doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Inches(0.28)
                            p.paragraph_format.line_spacing = 1.0
                            for run in p.runs:
                                set_font_style(run)
        
        elif text == '实验步骤':
            if i + 1 < len(template_doc.paragraphs):
                next_para = template_doc.paragraphs[i + 1]
                if '（具体内容' in next_para.text:
                    next_para.clear()
                    for line in content_sections.get('steps', []):
                        if line.strip():
                            p = template_doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Inches(0.28)
                            p.paragraph_format.line_spacing = 1.0
                            for run in p.runs:
                                set_font_style(run)
        
        elif text == '实验结果与分析':
            if i + 1 < len(template_doc.paragraphs):
                next_para = template_doc.paragraphs[i + 1]
                if '（具体内容' in next_para.text:
                    next_para.clear()
                    for line in content_sections.get('results', []):
                        if line.strip():
                            p = template_doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Inches(0.28)
                            p.paragraph_format.line_spacing = 1.0
                            for run in p.runs:
                                set_font_style(run)
        
        elif text == '实验总结':
            if i + 1 < len(template_doc.paragraphs):
                next_para = template_doc.paragraphs[i + 1]
                if '（具体内容' in next_para.text:
                    next_para.clear()
                    for line in content_sections.get('summary', []):
                        if line.strip():
                            p = template_doc.add_paragraph(line.strip())
                            p.paragraph_format.first_line_indent = Inches(0.28)
                            p.paragraph_format.line_spacing = 1.0
                            for run in p.runs:
                                set_font_style(run)
    
    # 保存
    template_doc.save(output_path)

# 尝试直接提取内容
try:
    # 直接读取你提供的两个文件
    content_sections = extract_sections_from_docx('f:/huibian/房价预测实验报告_模板样式.docx')
    
    # 填充模板
    fill_template_reliably('c:/Users/21268/Desktop/人工智能导论/人工智能导论实验报告模板.docx', 
                          content_sections, 
                          '房价预测实验报告_最终可靠版.docx')
    print("可靠填充完成！DOCX文件已生成：房价预测实验报告_最终可靠版.docx")
except Exception as e:
    print(f"可靠填充失败：{e}")
    # 如果失败，提供手动填充的备选方案
    print("如果自动填充失败，请确认文件路径正确，或尝试手动复制粘贴内容。")