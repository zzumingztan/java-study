import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT

def set_font_style(run, font_name='SimSun', font_size=Pt(10.5)):
    """设置字体样式（五号字体=10.5pt）"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size

def create_fixed_template_docx(md_content, output_path):
    # 创建新的文档
    doc = Document()
    
    # 设置页面为A4纸，适合双面打印
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.98)  # 2.5cm
        section.right_margin = Inches(0.98)  # 2.5cm
        section.top_margin = Inches(0.79)   # 2cm
        section.bottom_margin = Inches(0.79)  # 2cm
    
    # 设置默认字体为宋体五号（10.5pt）
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(10.5)
    
    # 按行处理Markdown内容
    lines = md_content.split('\n')
    
    # 添加郑州大学标题
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("郑州大学计算机与人工智能学院")
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run.font.size = Pt(16)
    run.bold = True
    p.add_run("\n\n")
    
    # 添加实验报告标题
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("人工智能实验报告")
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run.font.size = Pt(20)
    run.bold = True
    p.add_run("\n\n")
    
    # 添加实验题目
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("实验题目：     基于全连接网络实现房价预测                           ")
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run.font.size = Pt(16)
    run.bold = True
    p.add_run("\n\n")
    
    # 添加信息表格
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # 设置表格字体
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font_style(run)
    
    # 填充表格内容
    table.cell(0, 0).text = "专业班级"
    table.cell(0, 1).text = "人工智能六班"
    table.cell(1, 0).text = "学    号"
    table.cell(1, 1).text = "202415060613"
    table.cell(2, 0).text = "姓    名"
    table.cell(2, 1).text = "李润豪"
    table.cell(3, 0).text = "指导老师"
    table.cell(3, 1).text = "王可"
    table.cell(4, 0).text = "报告日期"
    table.cell(4, 1).text = "2026."
    
    # 添加空行
    doc.add_paragraph("\n")
    
    # 添加实验题目（重复）
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("实验题目：   基于全连接网络实现房价预测                                ")
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run.font.size = Pt(16)
    run.bold = True
    p.add_run("\n\n")
    
    # 处理正文内容 - 修复格式问题
    section_counter = 0
    section_titles = [
        "一、实验目的",
        "二、实验内容", 
        "三、实验环境",
        "四、实验步骤",
        "五、实验结果与分析",
        "六、实验总结"
    ]
    
    # 找到各个部分的起始位置
    content_lines = []
    in_main_content = False
    
    for line in lines:
        line = line.strip()
        if line.startswith('## 实验题目：') or line.startswith('---'):
            continue
        elif line.startswith('## 一、') or line.startswith('## 二、') or line.startswith('## 三、') or \
             line.startswith('## 四、') or line.startswith('## 五、') or line.startswith('## 六、'):
            # 跳过原始的标题行，我们用固定的标题
            continue
        elif line.startswith('一、') or line.startswith('二、') or line.startswith('三、') or \
             line.startswith('四、') or line.startswith('五、') or line.startswith('六、'):
            # 这是内容中的小标题，保留
            content_lines.append(line)
        elif line.startswith('- **硬件资源：**') or line.startswith('- **软件环境：**') or line.startswith('- **开发平台：**'):
            # 保留环境部分的列表项
            content_lines.append(line)
        elif line.startswith('### ') or line.startswith('#### ') or line.startswith('##### ') or line.startswith('###### '):
            # 跳过子标题
            continue
        elif line.startswith('```') or line.startswith('![') or line.startswith('[') or line.startswith('* '):
            # 跳过代码块、图片、链接、列表等
            continue
        elif line and not line.startswith('#') and not line.startswith('*') and not line.startswith('-'):
            # 普通段落，保留
            content_lines.append(line)
    
    # 添加固定格式的章节标题
    for i, title in enumerate(section_titles):
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        run = p.add_run(title)
        run.font.name = 'SimHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
        run.font.size = Pt(14)
        run.bold = True
        p.add_run("\n")
        
        # 添加对应章节的内容
        if i == 0:  # 实验目的
            # 找到实验目的部分
            start_idx = -1
            for j, line in enumerate(lines):
                if line.strip().startswith('一、实验目的'):
                    start_idx = j + 1
                    break
            if start_idx != -1:
                # 添加实验目的内容
                j = start_idx
                while j < len(lines) and not lines[j].strip().startswith('二、') and not lines[j].strip().startswith('## '):
                    line = lines[j].strip()
                    if line and not line.startswith('#') and not line.startswith('*') and not line.startswith('-'):
                        p = doc.add_paragraph(line)
                        p.paragraph_format.first_line_indent = Inches(0.28)
                        p.paragraph_format.line_spacing = 1.0
                        for run in p.runs:
                            set_font_style(run)
                    j += 1
        elif i == 1:  # 实验内容
            # 添加实验内容
            p = doc.add_paragraph("飞桨内部集成了波士顿房价数据集 uci-housing，该数据集共 506 条数据，每条数据包含 14 个值，其中前 13 个值用来描述房屋的各种信息，作为特征信息，最后一个值为该类房屋价格中位数，即要回归预测的值。")
            p.paragraph_format.first_line_indent = Inches(0.28)
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                set_font_style(run)
            
            p = doc.add_paragraph("本实验要求使用飞桨搭建单层的全连接网络，并利用上述数据集训练该模型；训练结束后保存模型，存储模型的参数；评估网络时，从参数文件中加载模型参数，再评估模型。")
            p.paragraph_format.first_line_indent = Inches(0.28)
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                set_font_style(run)
        elif i == 2:  # 实验环境
            p = doc.add_paragraph("- **硬件资源：** GPU")
            p.paragraph_format.first_line_indent = Inches(0.28)
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                set_font_style(run)
            
            p = doc.add_paragraph("- **软件环境：** Python 3+, PaddlePaddle 2.0.2, NumPy, Matplotlib")
            p.paragraph_format.first_line_indent = Inches(0.28)
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                set_font_style(run)
            
            p = doc.add_paragraph("- **开发平台：** Baidu AI Studio")
            p.paragraph_format.first_line_indent = Inches(0.28)
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                set_font_style(run)
        elif i == 3:  # 实验步骤
            # 添加实验步骤内容（简化）
            p = doc.add_paragraph("本实验包括以下步骤：数据加载及预处理、模型配置、模型训练、模型评估。")
            p.paragraph_format.first_line_indent = Inches(0.28)
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                set_font_style(run)
        elif i == 4:  # 实验结果与分析
            # 添加实验结果与分析
            p = doc.add_paragraph("通过训练损失曲线和预测值与真实值对比图分析，模型在200轮迭代后收敛到17~28区间，表明模型对房价数据具有良好的拟合能力。")
            p.paragraph_format.first_line_indent = Inches(0.28)
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                set_font_style(run)
        elif i == 5:  # 实验总结
            # 添加实验总结
            p = doc.add_paragraph("通过本次实验，我深入理解了全连接神经网络在回归任务中的应用，掌握了深度学习回归任务的标准流程，包括数据加载、模型定义、损失函数与优化器选择、循环训练、模型保存与评估等关键环节。")
            p.paragraph_format.first_line_indent = Inches(0.28)
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                set_font_style(run)
        
        # 添加空行
        doc.add_paragraph()
    
    # 保存文档
    doc.save(output_path)

# 读取Markdown文件
try:
    with open('房价预测实验报告.md', 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 转换为DOCX（修复格式）
    create_fixed_template_docx(md_content, '房价预测实验报告_最终版.docx')
    print("格式修复完成！DOCX文件已生成：房价预测实验报告_最终版.docx")
except Exception as e:
    print(f"格式修复失败：{e}")