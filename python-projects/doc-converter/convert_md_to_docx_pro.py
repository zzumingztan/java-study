import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT

def set_font_style(run, font_name='SimSun', font_size=Pt(10.5)):
    """设置字体样式（五号字体=10.5pt）"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size

def create_docx_from_md(md_content, output_path):
    # 创建新的文档
    doc = Document()
    
    # 设置页面为双面打印（A4纸，左右页边距2.5cm）
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.98)  # 2.5cm
        section.right_margin = Inches(0.98)  # 2.5cm
        section.top_margin = Inches(0.79)   # 2cm
        section.bottom_margin = Inches(0.79)  # 2cm
    
    # 设置默认字体为宋体五号（10.5pt）
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(10.5)
    
    # 按行处理Markdown内容
    lines = md_content.split('\n')
    
    # 标题级别映射
    heading_levels = {
        '# ': 0,      # 一级标题（居中，黑体，小二号）
        '## ': 1,     # 二级标题（左对齐，黑体，小三号）
        '### ': 2,    # 三级标题（左对齐，黑体，四号）
        '#### ': 3,   # 四级标题（左对齐，黑体，小四号）
        '##### ': 4,  # 五级标题（左对齐，宋体，小四号）
        '###### ': 5  # 六级标题（左对齐，宋体，五号）
    }
    
    for line in lines:
        line = line.strip()
        if not line:
            # 空行
            doc.add_paragraph()
            continue
            
        # 标题处理
        is_heading = False
        for prefix, level in heading_levels.items():
            if line.startswith(prefix):
                heading_text = line[len(prefix):]
                p = doc.add_heading(heading_text, level=level+1)
                
                # 设置标题字体
                if level == 0:  # 一级标题
                    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    for run in p.runs:
                        run.font.name = 'SimHei'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
                        run.font.size = Pt(18)  # 小二号
                elif level == 1:  # 二级标题
                    for run in p.runs:
                        run.font.name = 'SimHei'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
                        run.font.size = Pt(15)  # 小三号
                elif level == 2:  # 三级标题
                    for run in p.runs:
                        run.font.name = 'SimHei'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
                        run.font.size = Pt(14)  # 四号
                elif level == 3:  # 四级标题
                    for run in p.runs:
                        run.font.name = 'SimHei'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
                        run.font.size = Pt(12)  # 小四号
                else:  # 五级和六级标题
                    for run in p.runs:
                        run.font.name = 'SimSun'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimSun')
                        run.font.size = Pt(12)  # 小四号或五号
                
                is_heading = True
                break
        
        if is_heading:
            continue
            
        # 表格处理（简化处理）
        if '|' in line and line.strip().startswith('|') and line.strip().endswith('|'):
            p = doc.add_paragraph(line)
            for run in p.runs:
                set_font_style(run)
            continue
            
        # 代码块处理
        if line.startswith('```'):
            continue
            
        # 列表处理
        if line.startswith('- [ ]') or line.startswith('- [x]'):
            p = doc.add_paragraph(line, style='List Bullet')
            for run in p.runs:
                set_font_style(run)
            continue
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            for run in p.runs:
                set_font_style(run)
            continue
            
        # 强调和斜体处理
        if '**' in line:
            p = doc.add_paragraph()
            parts = line.split('**')
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    if part.strip():
                        run = p.add_run(part)
                        set_font_style(run)
                else:
                    if part.strip():
                        run = p.add_run(part)
                        set_font_style(run)
                        run.bold = True
            continue
        elif '*' in line and not line.startswith('* ') and not line.endswith(' *'):
            p = doc.add_paragraph()
            parts = line.split('*')
            for i, part in enumerate(parts):
                if i % 2 == 0:
                    if part.strip():
                        run = p.add_run(part)
                        set_font_style(run)
                else:
                    if part.strip():
                        run = p.add_run(part)
                        set_font_style(run)
                        run.italic = True
            continue
            
        # 普通段落（首行缩进2字符，单倍行距，五号字体）
        if line.startswith('一、') or line.startswith('二、') or line.startswith('三、') or \
           line.startswith('四、') or line.startswith('五、') or line.startswith('六、'):
            # 实验部分标题
            p = doc.add_paragraph(line)
            p.paragraph_format.first_line_indent = Inches(0.28)  # 2字符缩进约0.28英寸
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                set_font_style(run)
        else:
            # 普通段落
            p = doc.add_paragraph(line)
            p.paragraph_format.first_line_indent = Inches(0.28)  # 首行缩进2字符
            p.paragraph_format.line_spacing = 1.0  # 单倍行距
            for run in p.runs:
                set_font_style(run)
    
    # 保存文档
    doc.save(output_path)

# 读取Markdown文件
try:
    with open('房价预测实验报告.md', 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 转换为DOCX
    create_docx_from_md(md_content, '房价预测实验报告_格式化.docx')
    print("格式化转换成功！DOCX文件已生成：房价预测实验报告_格式化.docx")
except Exception as e:
    print(f"格式化转换失败：{e}")