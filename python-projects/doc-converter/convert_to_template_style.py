import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT

def set_font_style(run, font_name='SimSun', font_size=Pt(10.5)):
    """设置字体样式（五号字体=10.5pt）"""
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = font_size

def create_template_style_docx(md_content, output_path):
    # 创建新的文档
    doc = Document()
    
    # 设置页面为A4纸，适合双面打印
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.98)  # 2.5cm
        section.right_margin = Inches(0.98)  # 2.5cm
        section.top_margin = Inches(0.79)   # 2cm
        section.bottom_margin = Inches(0.79)  # 2cm
    
    # 设置默认字体为宋体五号（10.5pt）
    style = doc.styles['Normal']
    font = style.font
    font.name = 'SimSun'
    font.size = Pt(10.5)
    
    # 按行处理Markdown内容
    lines = md_content.split('\n')
    
    # 标题级别映射
    heading_levels = {
        '# ': 0,      # 一级标题（居中，黑体，小二号）
        '## ': 1,     # 二级标题（左对齐，黑体，小三号）
        '### ': 2,    # 三级标题（左对齐，黑体，四号）
        '#### ': 3,   # 四级标题（左对齐，黑体，小四号）
        '##### ': 4,  # 五级标题（左对齐，宋体，小四号）
        '###### ': 5  # 六级标题（左对齐，宋体，五号）
    }
    
    # 添加郑州大学标题
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("郑州大学计算机与人工智能学院")
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run.font.size = Pt(16)
    run.bold = True
    p.add_run("\n\n")
    
    # 添加实验报告标题
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("人工智能实验报告")
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run.font.size = Pt(20)
    run.bold = True
    p.add_run("\n\n")
    
    # 添加实验题目
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("实验题目：     基于全连接网络实现房价预测                           ")
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run.font.size = Pt(16)
    run.bold = True
    p.add_run("\n\n")
    
    # 添加信息表格
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Table Grid'
    
    # 设置表格字体
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    set_font_style(run)
    
    # 填充表格内容
    table.cell(0, 0).text = "专业班级"
    table.cell(0, 1).text = "人工智能六班"
    table.cell(1, 0).text = "学    号"
    table.cell(1, 1).text = "202415060613"
    table.cell(2, 0).text = "姓    名"
    table.cell(2, 1).text = "李润豪"
    table.cell(3, 0).text = "指导老师"
    table.cell(3, 1).text = "王可"
    table.cell(4, 0).text = "报告日期"
    table.cell(4, 1).text = "2026."
    
    # 添加空行
    doc.add_paragraph("\n")
    
    # 添加实验题目（重复）
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    run = p.add_run("实验题目：   基于全连接网络实现房价预测                                ")
    run.font.name = 'SimHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
    run.font.size = Pt(16)
    run.bold = True
    p.add_run("\n\n")
    
    # 处理正文内容
    in_section = False
    for line in lines:
        line = line.strip()
        if not line:
            # 空行
            doc.add_paragraph()
            continue
            
        # 检查是否是实验部分标题
        if line.startswith('## 一、') or line.startswith('## 二、') or line.startswith('## 三、') or \
           line.startswith('## 四、') or line.startswith('## 五、') or line.startswith('## 六、'):
            # 实验部分标题（加粗）
            section_title = line[4:]  # 去掉'## '
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = p.add_run(section_title)
            run.font.name = 'SimHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'SimHei')
            run.font.size = Pt(14)
            run.bold = True
            p.add_run("\n")
            in_section = True
            continue
            
        # 普通段落（首行缩进2字符，单倍行距，五号字体）
        if in_section and (line.startswith('一、') or line.startswith('二、') or line.startswith('三、') or \
                          line.startswith('四、') or line.startswith('五、') or line.startswith('六、')):
            # 实验部分标题
            p = doc.add_paragraph(line)
            p.paragraph_format.first_line_indent = Inches(0.28)  # 2字符缩进约0.28英寸
            p.paragraph_format.line_spacing = 1.0
            for run in p.runs:
                set_font_style(run)
        elif in_section:
            # 普通段落
            p = doc.add_paragraph(line)
            p.paragraph_format.first_line_indent = Inches(0.28)  # 首行缩进2字符
            p.paragraph_format.line_spacing = 1.0  # 单倍行距
            for run in p.runs:
                set_font_style(run)
        else:
            # 跳过非正文内容
            continue
    
    # 保存文档
    doc.save(output_path)

# 读取Markdown文件
try:
    with open('房价预测实验报告.md', 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 转换为DOCX（模板样式）
    create_template_style_docx(md_content, '房价预测实验报告_模板样式.docx')
    print("模板样式转换成功！DOCX文件已生成：房价预测实验报告_模板样式.docx")
except Exception as e:
    print(f"模板样式转换失败：{e}")