from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# 样式设置
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5

for i in range(1, 4):
    heading_style = doc.styles[f'Heading {i}']
    heading_font = heading_style.font
    heading_font.name = '黑体'
    heading_font.color.rgb = RGBColor(0, 0, 0)
    if i == 1:
        heading_font.size = Pt(16)
        heading_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif i == 2:
        heading_font.size = Pt(14)
    else:
        heading_font.size = Pt(12)

# 标题
doc.add_heading('Java Swing 复选框（JCheckBox）实验报告', level=1)

# 一、实验目的
doc.add_heading('一、实验目的', level=2)
for p_text in [
    '掌握 Java Swing 中 JCheckBox 复选框组件的使用方法。',
    '掌握 ActionListener 事件监听器的使用，实现对复选框选择状态变化的响应。',
    '理解 Swing 布局管理器（BorderLayout、FlowLayout）的基本用法。',
    '掌握 JTextField 文本框用于显示动态结果的技术。'
]:
    doc.add_paragraph(p_text, style='List Number')

# 二、实验环境
doc.add_heading('二、实验环境', level=2)
env_table = doc.add_table(rows=5, cols=2, style='Light Grid Accent 1')
for i, (k, v) in enumerate([
    ('操作系统', 'Windows 11'),
    ('JDK 版本', 'Java 8 及以上'),
    ('开发工具', 'VS Code / IntelliJ IDEA'),
    ('GUI 框架', 'Java Swing（javax.swing）'),
    ('运行工具', 'Eclipse / IDEA / 命令行')
]):
    env_table.rows[i].cells[0].text = k
    env_table.rows[i].cells[1].text = v

# 三、实验内容与要求
doc.add_heading('三、实验内容与要求', level=2)
doc.add_paragraph('设计并实现一个 Swing 图形用户界面程序，具体要求如下：')
for req in [
    '创建一个窗口，标题为"关于复选框"。',
    '窗口中包含三个复选框：音乐、运动、上网。',
    '当用户点击复选框时，在下方文本框中实时显示当前选中的爱好选项。',
    '窗口大小设置为 400×200 像素，启动时居中显示。'
]:
    doc.add_paragraph(req, style='List Number')

# 四、程序设计与实现
doc.add_heading('四、程序设计与实现', level=2)

doc.add_heading('4.1 整体架构', level=3)
doc.add_paragraph('程序包含一个类 CheckBoxDemo，继承自 JFrame，并实现 ActionListener 接口。类结构如下：')

member_table = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
for i, (a, b, c) in enumerate([
    ('成员', '类型', '说明'),
    ('cbMusic', 'JCheckBox', '"音乐"复选框'),
    ('cbSport', 'JCheckBox', '"运动"复选框'),
    ('cbInternet', 'JCheckBox', '"上网"复选框'),
    ('tfResult', 'JTextField', '显示选中结果的只读文本框'),
    ('label', 'JLabel', '"爱好选择："标签'),
]):
    member_table.rows[i].cells[0].text = a
    member_table.rows[i].cells[1].text = b
    member_table.rows[i].cells[2].text = c

doc.add_heading('4.2 界面布局', level=3)
doc.add_paragraph('采用嵌套面板的布局方式：')
doc.add_paragraph('• 外层面板（panel）：使用 BorderLayout，将窗口分为 NORTH 和 CENTER 两个区域。')
doc.add_paragraph('• 顶部面板（topPanel）：使用 FlowLayout（左对齐），放置标签和三个复选框。')
doc.add_paragraph('• 底部面板（bottomPanel）：使用 FlowLayout（左对齐），放置标签和结果显示文本框。')

doc.add_heading('4.3 核心代码分析', level=3)

doc.add_paragraph('（1）窗口初始化设置', style='List Bullet')
for line in [
    'setTitle("关于复选框");                              // 设置窗口标题',
    'setSize(400, 200);                                   // 设置窗口尺寸',
    'setDefaultCloseOperation(JFrame.EXIT_ON_CLOSE);      // 点击关闭按钮时退出程序',
    'setLocationRelativeTo(null);                         // 窗口在屏幕居中显示'
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(line)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

doc.add_paragraph('（2）组件创建与属性设置', style='List Bullet')
for line in [
    'cbMusic = new JCheckBox("音乐");',
    'cbSport = new JCheckBox("运动");',
    'cbInternet = new JCheckBox("上网");',
    'tfResult = new JTextField(20);',
    'tfResult.setEditable(false);   // 设置文本框为只读'
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(line)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

doc.add_paragraph('（3）事件监听注册', style='List Bullet')
doc.add_paragraph('三个复选框都注册了同一个 ActionListener，实现对点击事件的统一处理。')
for line in [
    'cbMusic.addActionListener(this);',
    'cbSport.addActionListener(this);',
    'cbInternet.addActionListener(this);'
]:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(line)
    run.font.name = 'Consolas'
    run.font.size = Pt(10)

doc.add_paragraph('（4）事件处理逻辑', style='List Bullet')
doc.add_paragraph('处理流程：创建 StringBuilder → 依次检查复选框选中状态 → 拼接文本 → trim() 去空格后更新文本框。')
event_code = (
    '@Override\n'
    'public void actionPerformed(ActionEvent e) {\n'
    '    StringBuilder sb = new StringBuilder();\n'
    '    if (cbMusic.isSelected()) { sb.append("音乐 "); }\n'
    '    if (cbSport.isSelected()) { sb.append("运动 "); }\n'
    '    if (cbInternet.isSelected()) { sb.append("上网 "); }\n'
    '    tfResult.setText(sb.toString().trim());\n'
    '}'
)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(event_code)
run.font.name = 'Consolas'
run.font.size = Pt(10)

doc.add_paragraph('（5）程序启动', style='List Bullet')
doc.add_paragraph('使用 SwingUtilities.invokeLater() 确保 GUI 创建在事件调度线程（EDT）中执行，保证线程安全。')
startup_code = (
    'public static void main(String[] args) {\n'
    '    SwingUtilities.invokeLater(() -> {\n'
    '        CheckBoxDemo demo = new CheckBoxDemo();\n'
    '        demo.setVisible(true);\n'
    '    });\n'
    '}'
)
p = doc.add_paragraph()
p.paragraph_format.left_indent = Cm(1)
run = p.add_run(startup_code)
run.font.name = 'Consolas'
run.font.size = Pt(10)

# 五、技术要点
doc.add_heading('五、技术要点', level=2)
doc.add_heading('5.1 Swing 线程安全', level=3)
doc.add_paragraph('所有 Swing GUI 操作必须在事件调度线程（EDT）中执行。使用 SwingUtilities.invokeLater() 将 GUI 创建代码放入 EDT 执行队列，避免线程安全问题。')
doc.add_heading('5.2 布局管理器组合', level=3)
doc.add_paragraph('本程序采用了 BorderLayout + FlowLayout 的组合布局方式，实现了清晰的面板区域划分。')
doc.add_heading('5.3 StringBuilder 优化字符串拼接', level=3)
doc.add_paragraph('在事件处理方法中使用 StringBuilder 而非直接用 + 拼接字符串，属于良好的编程习惯。')

# 六、运行结果
doc.add_heading('六、运行结果', level=2)
doc.add_heading('6.1 初始界面', level=3)
doc.add_paragraph('程序启动后，窗口居中显示，三个复选框均未被选中，文本框为空。')
doc.add_heading('6.2 交互测试', level=3)
result_table = doc.add_table(rows=6, cols=3, style='Light Grid Accent 1')
for i, (a, b, c) in enumerate([
    ('操作', '预期结果', '实际结果'),
    ('勾选"音乐"', '文本框显示"音乐"', '✅ 符合预期'),
    ('勾选"运动"', '文本框显示"音乐 运动"', '✅ 符合预期'),
    ('勾选"上网"', '文本框显示"音乐 运动 上网"', '✅ 符合预期'),
    ('取消"运动"', '文本框显示"音乐 上网"', '✅ 符合预期'),
    ('全部取消', '文本框为空', '✅ 符合预期'),
]):
    result_table.rows[i].cells[0].text = a
    result_table.rows[i].cells[1].text = b
    result_table.rows[i].cells[2].text = c

# 七、实验总结
doc.add_heading('七、实验总结', level=2)
for i, s in enumerate([
    'JCheckBox 复选框是 Swing 中用于多项选择的标准组件，通过 isSelected() 方法可以获取当前选中状态。',
    '复选框触发 ActionEvent 事件时，可通过实现 ActionListener 接口来处理选择状态变化。',
    '合理使用布局管理器组合（BorderLayout + FlowLayout）可以快速构建整洁的界面布局。',
    'SwingUtilities.invokeLater() 是确保 Swing 程序线程安全的标准方式。',
    '实验中涉及的知识点包括：事件监听模型、布局管理、组件属性设置、字符串处理等，是 Java GUI 编程的重要基础。',
], 1):
    doc.add_paragraph(f'{i}. {s}')
doc.add_paragraph()
doc.add_paragraph('通过本次实验，我掌握了 Swing 复选框组件的使用方法以及事件处理机制，为后续开发更复杂的 Java GUI 应用程序打下了基础。')

# 保存
save_path = r'F:\IDE\study\basic-code\src\work9\Java_Swing复选框实验报告.docx'
doc.save(save_path)
print(f'Word文档已保存到: {save_path}')
print(f'文件大小: {os.path.getsize(save_path)} bytes')
